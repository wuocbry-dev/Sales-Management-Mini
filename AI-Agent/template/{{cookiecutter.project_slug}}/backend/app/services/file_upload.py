{%- if cookiecutter.use_jwt and cookiecutter.use_postgresql %}
"""File upload service (PostgreSQL async).

Contains business logic for file validation, content parsing, and chat file
creation. Moves parsing helpers and file classification out of the route layer.
"""

import logging
from typing import Any
{%- if cookiecutter.use_all_pdf_parsers or cookiecutter.use_llamaparse %}
import os
import tempfile
{%- endif %}

from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.chat_file import ChatFile
from app.services.file_storage import (
    ALLOWED_MIME_TYPES,
    MAX_UPLOAD_SIZE,
    classify_file,
)

logger = logging.getLogger(__name__)


class FileUploadService:
    """Service for file upload validation, parsing, and persistence."""

    ALLOWED_MIME_TYPES = ALLOWED_MIME_TYPES
    MAX_UPLOAD_SIZE = MAX_UPLOAD_SIZE

    def __init__(self, db: AsyncSession):
        self.db = db

    @staticmethod
    def validate_upload(content_type: str | None, size: int) -> tuple[bool, str | None]:
        """Validate file type and size.

        Returns:
            Tuple of (is_valid, error_message).
        """
        if content_type not in ALLOWED_MIME_TYPES:
            return False, f"File type '{content_type}' is not supported."
        if size > MAX_UPLOAD_SIZE:
            return False, f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024 * 1024)}MB."
        return True, None

    @staticmethod
    def classify_file(mime_type: str, filename: str) -> str:
        """Classify file type based on MIME type and extension."""
        return classify_file(mime_type, filename)

    async def parse_content(
        self,
        data: bytes,
        file_type: str,
        mime_type: str = "",
    ) -> str | None:
        """Parse file content based on file type.

        Returns extracted text content or None if parsing fails.
        """
        if file_type == "text":
            return self._parse_text_content(data, mime_type)
{%- if cookiecutter.use_all_pdf_parsers %}
        elif file_type == "pdf":
            return await self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
{%- elif not cookiecutter.use_llamaparse %}
        elif file_type == "pdf":
            return self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
{%- endif %}
        return None

    @staticmethod
    def _parse_text_content(data: bytes, mime_type: str) -> str | None:
        """Extract text content from text-based files."""
        try:
            return data.decode("utf-8")
        except (UnicodeDecodeError, ValueError):
            return None

{%- if cookiecutter.use_all_pdf_parsers %}

    @staticmethod
    def _parse_pdf_pymupdf(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call]
            text_parts = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text.strip())
            doc.close()
            return "\n\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.warning(f"PyMuPDF PDF parsing failed: {e}")
            return None

    async def _parse_pdf_llamaparse(self, data: bytes) -> str | None:
        """Extract text from PDF using LlamaParse."""
        try:
            from llama_cloud import AsyncLlamaCloud
            from app.core.config import settings

            if not settings.LLAMAPARSE_API_KEY:
                logger.warning("LLAMAPARSE_API_KEY not set, falling back to PyMuPDF")
                return self._parse_pdf_pymupdf(data)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(data)
                temp_path = f.name
            try:
                client = AsyncLlamaCloud(api_key=settings.LLAMAPARSE_API_KEY)
                result = await client.parsing.upload_and_parse(
                    file=open(temp_path, "rb"),
                    tier=settings.LLAMAPARSE_TIER,
                )
                return "\n\n".join(p.markdown for p in result.pages) if result.pages else None
            finally:
                os.unlink(temp_path)
        except Exception as e:
            logger.warning(f"LlamaParse PDF parsing failed: {e}")
            return self._parse_pdf_pymupdf(data)

    async def _parse_pdf_liteparse(self, data: bytes) -> str | None:
        """Extract text from PDF using LiteParse."""
        try:
            from liteparse import LiteParse

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(data)
                temp_path = f.name
            try:
                parser = LiteParse()
                result = await parser.aparse(temp_path)
                pages = result.pages if hasattr(result, "pages") else [result]
                text = "\n\n".join(
                    p.content if hasattr(p, "content") else str(p) for p in pages
                )
                return text if text.strip() else None
            finally:
                os.unlink(temp_path)
        except Exception as e:
            logger.warning(f"LiteParse PDF parsing failed: {e}")
            return self._parse_pdf_pymupdf(data)

    async def _parse_pdf_content(self, data: bytes) -> str | None:
        """Parse PDF using the parser selected by CHAT_PDF_PARSER env var."""
        from app.core.config import settings

        parser = getattr(settings, "CHAT_PDF_PARSER", "pymupdf")
        if parser == "llamaparse":
            return await self._parse_pdf_llamaparse(data)
        elif parser == "liteparse":
            return await self._parse_pdf_liteparse(data)
        return self._parse_pdf_pymupdf(data)

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io
            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None

{%- elif not cookiecutter.use_llamaparse %}

    @staticmethod
    def _parse_pdf_content(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call,unused-ignore]
            texts = []
            for page in doc:
                blocks = page.get_text("blocks")
                for b in blocks:
                    if b[6] == 0:
                        text = b[4].strip()
                        if text:
                            texts.append(text)
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for table in tables.tables:
                            df = table.to_pandas()
                            if not df.empty:
                                texts.append(df.to_markdown(index=False))
                except Exception:
                    pass
            doc.close()
            return "\n\n".join(texts) if texts else None
        except Exception as e:
            logger.warning(f"PDF parsing failed: {e}")
            return None

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io
            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None
{%- endif %}

    async def get_user_file(self, file_id: Any, user_id: Any) -> ChatFile:
        """Get a file by ID, verifying ownership.

        Raises:
            NotFoundError: If file does not exist or user has no access.
        """
        from app.core.exceptions import NotFoundError
        from app.repositories import chat_file as chat_file_repo

        chat_file = await chat_file_repo.get_by_id(self.db, file_id)
        if not chat_file or str(chat_file.user_id) != str(user_id):
            raise NotFoundError(message="File not found")
        return chat_file

    async def create_chat_file(
        self,
        *,
        user_id: Any,
        filename: str,
        mime_type: str,
        size: int,
        storage_path: str,
        file_type: str,
        parsed_content: str | None = None,
    ) -> ChatFile:
        """Create a chat file record in the database."""
        chat_file = ChatFile(
            user_id=user_id,
            filename=filename,
            mime_type=mime_type,
            size=size,
            storage_path=storage_path,
            file_type=file_type,
            parsed_content=parsed_content,
        )
        self.db.add(chat_file)
        await self.db.flush()
        await self.db.commit()
        await self.db.refresh(chat_file)
        return chat_file


{%- elif cookiecutter.use_jwt and cookiecutter.use_sqlite %}
"""File upload service (SQLite sync).

Contains business logic for file validation, content parsing, and chat file
creation. Moves parsing helpers and file classification out of the route layer.
"""

import logging
from typing import Any
{%- if cookiecutter.use_all_pdf_parsers or cookiecutter.use_llamaparse %}
import os
import tempfile
{%- endif %}

from sqlalchemy.orm import Session

from app.db.models.chat_file import ChatFile
from app.services.file_storage import (
    ALLOWED_MIME_TYPES,
    MAX_UPLOAD_SIZE,
    classify_file,
)

logger = logging.getLogger(__name__)


class FileUploadService:
    """Service for file upload validation, parsing, and persistence."""

    ALLOWED_MIME_TYPES = ALLOWED_MIME_TYPES
    MAX_UPLOAD_SIZE = MAX_UPLOAD_SIZE

    def __init__(self, db: Session):
        self.db = db

    @staticmethod
    def validate_upload(content_type: str | None, size: int) -> tuple[bool, str | None]:
        """Validate file type and size.

        Returns:
            Tuple of (is_valid, error_message).
        """
        if content_type not in ALLOWED_MIME_TYPES:
            return False, f"File type '{content_type}' is not supported."
        if size > MAX_UPLOAD_SIZE:
            return False, f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024 * 1024)}MB."
        return True, None

    @staticmethod
    def classify_file(mime_type: str, filename: str) -> str:
        """Classify file type based on MIME type and extension."""
        return classify_file(mime_type, filename)

    def parse_content(
        self,
        data: bytes,
        file_type: str,
        mime_type: str = "",
    ) -> str | None:
        """Parse file content based on file type.

        Returns extracted text content or None if parsing fails.
        """
        if file_type == "text":
            return self._parse_text_content(data, mime_type)
{%- if cookiecutter.use_all_pdf_parsers %}
        elif file_type == "pdf":
            return self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
{%- elif not cookiecutter.use_llamaparse %}
        elif file_type == "pdf":
            return self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
{%- endif %}
        return None

    @staticmethod
    def _parse_text_content(data: bytes, mime_type: str) -> str | None:
        """Extract text content from text-based files."""
        try:
            return data.decode("utf-8")
        except (UnicodeDecodeError, ValueError):
            return None

{%- if cookiecutter.use_all_pdf_parsers %}

    @staticmethod
    def _parse_pdf_pymupdf(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call]
            text_parts = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text.strip())
            doc.close()
            return "\n\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.warning(f"PyMuPDF PDF parsing failed: {e}")
            return None

    def _parse_pdf_llamaparse(self, data: bytes) -> str | None:
        """Extract text from PDF using LlamaParse."""
        try:
            import asyncio
            from app.core.config import settings

            if not settings.LLAMAPARSE_API_KEY:
                logger.warning("LLAMAPARSE_API_KEY not set, falling back to PyMuPDF")
                return self._parse_pdf_pymupdf(data)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(data)
                temp_path = f.name
            try:
                async def _parse():
                    from llama_cloud import AsyncLlamaCloud
                    client = AsyncLlamaCloud(api_key=settings.LLAMAPARSE_API_KEY)
                    result = await client.parsing.upload_and_parse(
                        file=open(temp_path, "rb"),
                        tier=settings.LLAMAPARSE_TIER,
                    )
                    return "\n\n".join(p.markdown for p in result.pages) if result.pages else None

                return asyncio.run(_parse())
            finally:
                os.unlink(temp_path)
        except Exception as e:
            logger.warning(f"LlamaParse PDF parsing failed: {e}")
            return self._parse_pdf_pymupdf(data) if hasattr(self, '_parse_pdf_pymupdf') else None

    def _parse_pdf_liteparse(self, data: bytes) -> str | None:
        """Extract text from PDF using LiteParse (sync)."""
        try:
            from liteparse import LiteParse

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(data)
                temp_path = f.name
            try:
                parser = LiteParse()
                result = parser.parse(temp_path)
                pages = result.pages if hasattr(result, "pages") else [result]
                text = "\n\n".join(
                    p.content if hasattr(p, "content") else str(p) for p in pages
                )
                return text if text.strip() else None
            finally:
                os.unlink(temp_path)
        except Exception as e:
            logger.warning(f"LiteParse PDF parsing failed: {e}")
            return self._parse_pdf_pymupdf(data)

    def _parse_pdf_content(self, data: bytes) -> str | None:
        """Parse PDF using the parser selected by CHAT_PDF_PARSER env var."""
        from app.core.config import settings

        parser = getattr(settings, "CHAT_PDF_PARSER", "pymupdf")
        if parser == "llamaparse":
            return self._parse_pdf_llamaparse(data)
        elif parser == "liteparse":
            return self._parse_pdf_liteparse(data)
        return self._parse_pdf_pymupdf(data)

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io
            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None

{%- elif not cookiecutter.use_llamaparse %}

    @staticmethod
    def _parse_pdf_content(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call,unused-ignore]
            texts = []
            for page in doc:
                blocks = page.get_text("blocks")
                for b in blocks:
                    if b[6] == 0:
                        text = b[4].strip()
                        if text:
                            texts.append(text)
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for table in tables.tables:
                            df = table.to_pandas()
                            if not df.empty:
                                texts.append(df.to_markdown(index=False))
                except Exception:
                    pass
            doc.close()
            return "\n\n".join(texts) if texts else None
        except Exception as e:
            logger.warning(f"PDF parsing failed: {e}")
            return None

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io
            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None
{%- endif %}

    def get_user_file(self, file_id: Any, user_id: Any) -> ChatFile:
        """Get a file by ID, verifying ownership.

        Raises:
            NotFoundError: If file does not exist or user has no access.
        """
        from app.core.exceptions import NotFoundError
        from app.repositories import chat_file as chat_file_repo

        chat_file = chat_file_repo.get_by_id(self.db, file_id)
        if not chat_file or str(chat_file.user_id) != str(user_id):
            raise NotFoundError(message="File not found")
        return chat_file

    def create_chat_file(
        self,
        *,
        user_id: Any,
        filename: str,
        mime_type: str,
        size: int,
        storage_path: str,
        file_type: str,
        parsed_content: str | None = None,
    ) -> ChatFile:
        """Create a chat file record in the database."""
        chat_file = ChatFile(
            user_id=user_id,
            filename=filename,
            mime_type=mime_type,
            size=size,
            storage_path=storage_path,
            file_type=file_type,
            parsed_content=parsed_content,
        )
        self.db.add(chat_file)
        self.db.flush()
        self.db.commit()
        self.db.refresh(chat_file)
        return chat_file


{%- else %}
"""File upload service - not configured."""
{%- endif %}
