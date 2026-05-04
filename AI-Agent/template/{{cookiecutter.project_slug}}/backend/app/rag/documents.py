{%- if cookiecutter.enable_rag %}
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from langchain_text_splitters import RecursiveCharacterTextSplitter
{%- if cookiecutter.use_all_pdf_parsers %}
import pymupdf
from docx import Document as DOCXDocument
{%- elif cookiecutter.use_llamaparse %}
from llama_cloud import AsyncLlamaCloud
{%- elif cookiecutter.use_liteparse %}
from liteparse import LiteParse
from docx import Document as DOCXDocument
{%- else %}
import pymupdf
from docx import Document as DOCXDocument
{%- endif %}

from app.rag.config import RAGSettings, DocumentExtensions
from app.rag.models import Document, DocumentMetadata, DocumentPage, DocumentPageChunk
{%- if cookiecutter.enable_rag_image_description %}
from app.rag.models import DocumentImage
{%- endif %}

logger = logging.getLogger(__name__)


class BaseDocumentParser(ABC):
    """Abstract base class for document parsing strategies.
    Defines the interface that all document parsers must implement.
    Supports parsing of various document formats (PDF, DOCX, TXT, MD).
    """

    allowed = [f"{ext.value}" for ext in DocumentExtensions]

    def is_file_existing(self, filepath: Path) -> bool:
        """Check if file exists at the given path.
        Args:
            filepath: Path to the file to check.
        Returns:
            True if the file exists, False otherwise.
        """
        return Path.exists(filepath)

    def is_extension_allowed(self, filepath: Path) -> bool:
        """Check whether document extension is allowed for parsing.
        Args:
            filepath: Path to the file to check.
        Returns:
            True if the extension is supported and file exists.
        """
        return filepath.suffix.lower() in self.allowed and self.is_file_existing(filepath)

    def get_document_metadata(self, filepath: Path) -> DocumentMetadata:
        """Collect metadata about a given document.
        Args:
            filepath: Path to the document file.
        Returns:
            DocumentMetadata object containing file information.
        """
        import hashlib
        content_hash = hashlib.sha256(filepath.read_bytes()).hexdigest()
        return DocumentMetadata(
            filename=filepath.name,
            filesize=filepath.stat().st_size,
            filetype=filepath.suffix.replace(".", ""),
            source_path=str(filepath.resolve()),
            content_hash=content_hash,
        )

    @abstractmethod
    async def parse(self, filepath: Path) -> Document:
        """Parse a file and read its content into a Document object.
        Args:
            filepath: Path to the file to parse.
        Returns:
            Document object with parsed content and metadata.
        """
        pass


class TextDocumentParser(BaseDocumentParser):
    """Parser for text-based documents (TXT, MD).
    Uses Python's built-in file reading capabilities to extract
    text content from plain text and Markdown files.
    """

    def _parse_text_file(self, filepath: Path) -> Document:
        """Extract raw text from a TXT or MD file.
        Args:
            filepath: Path to the text file.
        Returns:
            Document object with the file content.
        """
        with open(filepath, "r", encoding="utf-8") as f:
            page = DocumentPage(
                page_num=1,
                content=f.read()
            )

        return Document(
            pages=[page],
            metadata=self.get_document_metadata(filepath)
        )

    async def parse(self, filepath: Path) -> Document:
        """Parse a text file (TXT or MD).

        Args:
            filepath: Path to the text file.

        Returns:
            Document object with parsed content.

        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by TextDocumentParser")

        if filepath.suffix in (".txt", ".md"):
            return self._parse_text_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


{%- if cookiecutter.use_all_pdf_parsers %}

class DocxDocumentParser(BaseDocumentParser):
    """Parser for DOCX documents using python-docx.

    Extracts text content from Microsoft Word documents by reading
    all paragraphs and joining them with newline characters.
    """

    def _parse_docx_file(self, filepath: Path) -> Document:
        """Extract raw text from the DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with the file content.
        """
        file: Any = DOCXDocument(str(filepath))
        page = DocumentPage(
            page_num=1,
            content="\n".join([p.text for p in file.paragraphs])
        )
        return Document(
            pages=[page],
            metadata=self.get_document_metadata(filepath)
        )

    async def parse(self, filepath: Path) -> Document:
        """Parse a DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with parsed content.

        Raises:
            ValueError: If the file is not a DOCX file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by DocxDocumentParser")

        if filepath.suffix == ".docx":
            return self._parse_docx_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


class PyMuPDFParser(BaseDocumentParser):
    """Smart PDF parser using PyMuPDF.

    Features:
    - Text extraction with layout preservation (blocks)
    - Table detection -> markdown tables
    - Header/footer detection and removal
    - OCR fallback for scanned pages (optional, requires tesseract)
    - Image extraction for LLM-based description
    - Document metadata (author, title, TOC)
    """

    MIN_TEXT_LENGTH = 50  # below this -> likely a scan, try OCR

    def __init__(self, enable_ocr: bool = False, image_describer: Any = None):
        self.enable_ocr = enable_ocr
        self._image_describer = image_describer

    def _detect_repeated_content(self, doc: Any) -> set[str]:
        """Detect headers/footers -- text appearing on >70% of pages."""
        if len(doc) < 3:
            return set()
        text_counts: dict[str, int] = {}
        for page in doc:
            for b in page.get_text("blocks"):
                if b[6] != 0:  # skip image blocks
                    continue
                y_ratio = b[1] / page.rect.height if page.rect.height else 0
                if y_ratio < 0.15 or y_ratio > 0.85:
                    text = b[4].strip()
                    if text and len(text) < 200:
                        text_counts[text] = text_counts.get(text, 0) + 1
        threshold = len(doc) * 0.7
        return {t for t, c in text_counts.items() if c >= threshold}

    def _extract_text(self, page: Any, repeated: set[str]) -> str:
        """Extract text blocks, filtering headers/footers."""
        texts = []
        for b in page.get_text("blocks"):
            if b[6] != 0:  # skip image blocks
                continue
            text = b[4].strip()
            if text and text not in repeated:
                texts.append(text)
        return str("\n\n".join(texts))

    def _extract_tables(self, page: Any) -> str:
        """Extract tables as markdown."""
        try:
            tables = page.find_tables()
            if not tables or not tables.tables:
                return ""
            parts = []
            for table in tables.tables:
                df = table.to_pandas()
                if not df.empty:
                    parts.append(df.to_markdown(index=False))
            return "\n\n".join(parts)
        except Exception:
            return ""

    def _ocr_page(self, page: Any, image_describer: Any = None) -> str:
        """OCR a scanned page by rendering it as image and sending to LLM vision."""
        if not image_describer:
            return ""
        try:
            import asyncio
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")
            loop = asyncio.new_event_loop()
            try:
                return str(loop.run_until_complete(
                    image_describer.describe(image_bytes, "image/png")
                ))
            finally:
                loop.close()
        except Exception as e:
            logger.warning(f"LLM OCR failed for page {page.number + 1}: {e}")
            return ""

{%- if cookiecutter.enable_rag_image_description %}
    def _extract_images(self, doc: Any, page: Any) -> list["DocumentImage"]:
        """Extract images from page for LLM description."""
        images = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
                if base and base["image"] and len(base["image"]) > 1000:
                    ext = base.get("ext", "png")
                    mime_map = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg"}
                    images.append(DocumentImage(
                        page_num=page.number + 1,
                        image_bytes=base["image"],
                        mime_type=mime_map.get(ext, f"image/{ext}"),
                    ))
            except Exception:
                pass
        return images
{%- endif %}

    def _parse_pdf_file(self, filepath: Path) -> Document:
        """Parse PDF with smart extraction pipeline."""
        doc: Any = pymupdf.open(filepath)  # type: ignore[no-untyped-call]

        # Doc-level metadata
        meta = doc.metadata or {}
        toc = doc.get_toc()

        # Detect repeated headers/footers
        repeated = self._detect_repeated_content(doc)

        pages = []
        for page in doc:
            # 1. Text with layout (skip image blocks, filter headers/footers)
            text = self._extract_text(page, repeated)

            # 2. Tables -> markdown
            tables_md = self._extract_tables(page)
            if tables_md:
                text = text + "\n\n" + tables_md if text.strip() else tables_md

            # 3. OCR fallback for scans/empty pages
            if self.enable_ocr and len(text.strip()) < self.MIN_TEXT_LENGTH:
                ocr_text = self._ocr_page(page, self._image_describer)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    logger.info(f"OCR fallback used for page {page.number + 1}")

{%- if cookiecutter.enable_rag_image_description %}
            # 4. Images
            images = self._extract_images(doc, page)
{%- endif %}

            pages.append(DocumentPage(
                page_num=page.number + 1,
                content=text,
{%- if cookiecutter.enable_rag_image_description %}
                images=images,
{%- endif %}
            ))

        doc.close()

        # Enrich metadata
        additional: dict[str, Any] = {}
        if meta.get("title"):
            additional["pdf_title"] = meta["title"]
        if meta.get("author"):
            additional["pdf_author"] = meta["author"]
        if toc:
            additional["toc"] = [{"level": t[0], "title": t[1], "page": t[2]} for t in toc[:20]]

        doc_meta = self.get_document_metadata(filepath)
        if additional:
            doc_meta.additional_info = {**(doc_meta.additional_info or {}), **additional}

        return Document(pages=pages, metadata=doc_meta)

    async def parse(self, filepath: Path) -> Document:
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by PyMuPDFParser")
        if filepath.suffix == ".pdf":
            return self._parse_pdf_file(filepath)
        raise ValueError(f"Unsupported: {filepath.suffix}")


class LlamaParseParser(BaseDocumentParser):
    """Advanced document parser using LlamaParse cloud API.

    Provides AI-powered document parsing with support for 130+ formats
    including PDF, DOCX, PPTX, XLSX, images (OCR), and more.
    Returns markdown-formatted content.
    """

    # LlamaParse supports these beyond our default allowed list
    EXTRA_SUPPORTED = {".pptx", ".xlsx", ".xls", ".csv", ".rtf", ".epub",
                       ".jpg", ".jpeg", ".png", ".html", ".htm"}

    def __init__(self, api_key: str, tier: str = "agentic"):
        """Initialize the LlamaParse parser.

        Args:
            api_key: LlamaCloud API key for authentication.
            tier: Parsing tier (fast, cost_effective, agentic, agentic_plus).
        """
        from llama_cloud import AsyncLlamaCloud
        self.parser = AsyncLlamaCloud(api_key=api_key)
        self.tier = tier
        # Extend allowed extensions with LlamaParse-supported formats
        self.allowed = [ext.value for ext in DocumentExtensions] + list(self.EXTRA_SUPPORTED)

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LlamaParse.

        Supports PDF, DOCX, PPTX, XLSX, images, and many more formats.
        See https://developers.llamaindex.ai/python/cloud/llamaparse/supported_document_types/

        Args:
            filepath: Path to the file to parse.

        Returns:
            Document object with parsed markdown content.

        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by LlamaParse")

        file_obj = await self.parser.files.create(file=filepath, purpose="parse")
        result = await self.parser.parsing.parse(
            file_id=file_obj.id,
            tier=self.tier,
            version="latest",
            expand=["text", "markdown"],
        )
        pages = []
        for page in result.markdown.pages:
            pages.append(DocumentPage(
                page_num=page.page_number,
                content=page.markdown
            ))

        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath)
        )


class LiteParseParser(BaseDocumentParser):
    """Document parser using LiteParse -- fast, local, AI-native parsing.

    Uses LiteParse (from LlamaIndex) for layout-aware text extraction.
    Preserves spatial relationships (tables as ASCII grids) instead of
    converting to markdown. Built-in OCR via Tesseract.js for scanned pages.

    Requires: pip install liteparse && npm i -g @llamaindex/liteparse
    """

    def __init__(self) -> None:
        from liteparse import LiteParse
        self.parser = LiteParse()

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LiteParse.

        LiteParse returns layout-aware text. If the result has per-page data
        it's used directly; otherwise full text is split on form-feed chars.
        """
        import asyncio

        # LiteParse Python wrapper is synchronous -- run in thread
        result = await asyncio.to_thread(self.parser.parse, str(filepath))

        pages: list[DocumentPage] = []

        # Try per-page output first
        if hasattr(result, "pages") and result.pages:
            for i, page in enumerate(result.pages):
                text = page.text if hasattr(page, "text") else str(page)
                if text.strip():
                    pages.append(DocumentPage(page_num=i + 1, content=text))
        else:
            # Fallback: split full text on form-feed (\f) as page separator
            full_text = result.text if hasattr(result, "text") else str(result)
            page_texts = full_text.split("\f") if "\f" in full_text else [full_text]
            for i, text in enumerate(page_texts):
                if text.strip():
                    pages.append(DocumentPage(page_num=i + 1, content=text.strip()))

        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath),
        )


class PdfParserFactory:
    """Factory for runtime PDF parser selection via PDF_PARSER env var."""

    @staticmethod
    def create(parser_name: str, settings: RAGSettings | None = None, image_describer: Any = None) -> BaseDocumentParser:
        if parser_name == "llamaparse":
            if not settings or not settings.pdf_parser.api_key:
                raise ValueError("LlamaParse requires LLAMAPARSE_API_KEY to be set")
            return LlamaParseParser(
                api_key=settings.pdf_parser.api_key,
                tier=settings.pdf_parser.tier,
            )
        elif parser_name == "liteparse":
            return LiteParseParser()
        else:
            return PyMuPDFParser(
                enable_ocr=settings.enable_ocr if settings else False,
                image_describer=image_describer,
            )

{%- elif not cookiecutter.use_llamaparse %}

class DocxDocumentParser(BaseDocumentParser):
    """Parser for DOCX documents using python-docx.

    Extracts text content from Microsoft Word documents by reading
    all paragraphs and joining them with newline characters.
    """

    def _parse_docx_file(self, filepath: Path) -> Document:
        """Extract raw text from the DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with the file content.
        """
        file: Any = DOCXDocument(str(filepath))
        page = DocumentPage(
            page_num=1,
            content="\n".join([p.text for p in file.paragraphs])
        )
        return Document(
            pages=[page],
            metadata=self.get_document_metadata(filepath)
        )

    async def parse(self, filepath: Path) -> Document:
        """Parse a DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with parsed content.

        Raises:
            ValueError: If the file is not a DOCX file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by DocxDocumentParser")

        if filepath.suffix == ".docx":
            return self._parse_docx_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


class PyMuPDFParser(BaseDocumentParser):
    """Smart PDF parser using PyMuPDF.

    Features:
    - Text extraction with layout preservation (blocks)
    - Table detection → markdown tables
    - Header/footer detection and removal
    - OCR fallback for scanned pages (optional, requires tesseract)
    - Image extraction for LLM-based description
    - Document metadata (author, title, TOC)
    """

    MIN_TEXT_LENGTH = 50  # below this → likely a scan, try OCR

    def __init__(self, enable_ocr: bool = False, image_describer: Any = None):
        self.enable_ocr = enable_ocr
        self._image_describer = image_describer

    def _detect_repeated_content(self, doc: Any) -> set[str]:
        """Detect headers/footers — text appearing on >70% of pages."""
        if len(doc) < 3:
            return set()
        text_counts: dict[str, int] = {}
        for page in doc:
            for b in page.get_text("blocks"):
                if b[6] != 0:  # skip image blocks
                    continue
                y_ratio = b[1] / page.rect.height if page.rect.height else 0
                if y_ratio < 0.15 or y_ratio > 0.85:
                    text = b[4].strip()
                    if text and len(text) < 200:
                        text_counts[text] = text_counts.get(text, 0) + 1
        threshold = len(doc) * 0.7
        return {t for t, c in text_counts.items() if c >= threshold}

    def _extract_text(self, page: Any, repeated: set[str]) -> str:
        """Extract text blocks, filtering headers/footers."""
        texts = []
        for b in page.get_text("blocks"):
            if b[6] != 0:  # skip image blocks
                continue
            text = b[4].strip()
            if text and text not in repeated:
                texts.append(text)
        return str("\n\n".join(texts))

    def _extract_tables(self, page: Any) -> str:
        """Extract tables as markdown."""
        try:
            tables = page.find_tables()
            if not tables or not tables.tables:
                return ""
            parts = []
            for table in tables.tables:
                df = table.to_pandas()
                if not df.empty:
                    parts.append(df.to_markdown(index=False))
            return "\n\n".join(parts)
        except Exception:
            return ""

    def _ocr_page(self, page: Any, image_describer: Any = None) -> str:
        """OCR a scanned page by rendering it as image and sending to LLM vision."""
        if not image_describer:
            return ""
        try:
            import asyncio
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")
            loop = asyncio.new_event_loop()
            try:
                return str(loop.run_until_complete(
                    image_describer.describe(image_bytes, "image/png")
                ))
            finally:
                loop.close()
        except Exception as e:
            logger.warning(f"LLM OCR failed for page {page.number + 1}: {e}")
            return ""

{%- if cookiecutter.enable_rag_image_description %}
    def _extract_images(self, doc: Any, page: Any) -> list["DocumentImage"]:
        """Extract images from page for LLM description."""
        images = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
                if base and base["image"] and len(base["image"]) > 1000:
                    ext = base.get("ext", "png")
                    mime_map = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg"}
                    images.append(DocumentImage(
                        page_num=page.number + 1,
                        image_bytes=base["image"],
                        mime_type=mime_map.get(ext, f"image/{ext}"),
                    ))
            except Exception:
                pass
        return images
{%- endif %}

    def _parse_pdf_file(self, filepath: Path) -> Document:
        """Parse PDF with smart extraction pipeline."""
        doc: Any = pymupdf.open(filepath)  # type: ignore[no-untyped-call]

        # Doc-level metadata
        meta = doc.metadata or {}
        toc = doc.get_toc()

        # Detect repeated headers/footers
        repeated = self._detect_repeated_content(doc)

        pages = []
        for page in doc:
            # 1. Text with layout (skip image blocks, filter headers/footers)
            text = self._extract_text(page, repeated)

            # 2. Tables → markdown
            tables_md = self._extract_tables(page)
            if tables_md:
                text = text + "\n\n" + tables_md if text.strip() else tables_md

            # 3. OCR fallback for scans/empty pages
            if self.enable_ocr and len(text.strip()) < self.MIN_TEXT_LENGTH:
                ocr_text = self._ocr_page(page, self._image_describer)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    logger.info(f"OCR fallback used for page {page.number + 1}")

{%- if cookiecutter.enable_rag_image_description %}
            # 4. Images
            images = self._extract_images(doc, page)
{%- endif %}

            pages.append(DocumentPage(
                page_num=page.number + 1,
                content=text,
{%- if cookiecutter.enable_rag_image_description %}
                images=images,
{%- endif %}
            ))

        doc.close()

        # Enrich metadata
        additional: dict[str, Any] = {}
        if meta.get("title"):
            additional["pdf_title"] = meta["title"]
        if meta.get("author"):
            additional["pdf_author"] = meta["author"]
        if toc:
            additional["toc"] = [{"level": t[0], "title": t[1], "page": t[2]} for t in toc[:20]]

        doc_meta = self.get_document_metadata(filepath)
        if additional:
            doc_meta.additional_info = {**(doc_meta.additional_info or {}), **additional}

        return Document(pages=pages, metadata=doc_meta)

    async def parse(self, filepath: Path) -> Document:
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by PyMuPDFParser")
        if filepath.suffix == ".pdf":
            return self._parse_pdf_file(filepath)
        raise ValueError(f"Unsupported: {filepath.suffix}")
{%- else %}

class LlamaParseParser(BaseDocumentParser):
    """Advanced document parser using LlamaParse cloud API.

    Provides AI-powered document parsing with support for 130+ formats
    including PDF, DOCX, PPTX, XLSX, images (OCR), and more.
    Returns markdown-formatted content.
    """

    # LlamaParse supports these beyond our default allowed list
    EXTRA_SUPPORTED = {".pptx", ".xlsx", ".xls", ".csv", ".rtf", ".epub",
                       ".jpg", ".jpeg", ".png", ".html", ".htm"}

    def __init__(self, api_key: str, tier: str = "agentic"):
        """Initialize the LlamaParse parser.

        Args:
            api_key: LlamaCloud API key for authentication.
            tier: Parsing tier (fast, cost_effective, agentic, agentic_plus).
        """
        self.parser = AsyncLlamaCloud(api_key=api_key)
        self.tier = tier
        # Extend allowed extensions with LlamaParse-supported formats
        self.allowed = [ext.value for ext in DocumentExtensions] + list(self.EXTRA_SUPPORTED)

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LlamaParse.

        Supports PDF, DOCX, PPTX, XLSX, images, and many more formats.
        See https://developers.llamaindex.ai/python/cloud/llamaparse/supported_document_types/

        Args:
            filepath: Path to the file to parse.

        Returns:
            Document object with parsed markdown content.

        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by LlamaParse")

        file_obj = await self.parser.files.create(file=filepath, purpose="parse")
        result = await self.parser.parsing.parse(
            file_id=file_obj.id,
            tier=self.tier,
            version="latest",
            expand=["text", "markdown"],
        )
        pages = []
        for page in result.markdown.pages:
            pages.append(DocumentPage(
                page_num=page.page_number,
                content=page.markdown
            ))

        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath)
        )
{%- endif %}

{%- if cookiecutter.use_liteparse and not cookiecutter.use_all_pdf_parsers %}


class LiteParseParser(BaseDocumentParser):
    """Document parser using LiteParse — fast, local, AI-native parsing.

    Uses LiteParse (from LlamaIndex) for layout-aware text extraction.
    Preserves spatial relationships (tables as ASCII grids) instead of
    converting to markdown. Built-in OCR via Tesseract.js for scanned pages.

    Requires: pip install liteparse && npm i -g @llamaindex/liteparse
    """

    def __init__(self) -> None:
        self.parser = LiteParse()

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LiteParse.

        LiteParse returns layout-aware text. If the result has per-page data
        it's used directly; otherwise full text is split on form-feed chars.
        """
        import asyncio

        # LiteParse Python wrapper is synchronous — run in thread
        result = await asyncio.to_thread(self.parser.parse, str(filepath))

        pages: list[DocumentPage] = []

        # Try per-page output first
        if hasattr(result, "pages") and result.pages:
            for i, page in enumerate(result.pages):
                text = page.text if hasattr(page, "text") else str(page)
                if text.strip():
                    pages.append(DocumentPage(page_num=i + 1, content=text))
        else:
            # Fallback: split full text on form-feed (\f) as page separator
            full_text = result.text if hasattr(result, "text") else str(result)
            page_texts = full_text.split("\f") if "\f" in full_text else [full_text]
            for i, text in enumerate(page_texts):
                if text.strip():
                    pages.append(DocumentPage(page_num=i + 1, content=text.strip()))

        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath),
        )
{%- endif %}


class DocumentProcessor:
    """Orchestrates parsing and chunking of files into Document objects.

    Manages the document processing pipeline:
    1. Route to appropriate parser based on file extension
    2. Parse document content
    3. Chunk document pages using RecursiveCharacterTextSplitter

{%- if cookiecutter.use_all_pdf_parsers %}
    Supported file types:
    - TXT, MD: TextDocumentParser (Python native)
    - DOCX: DocxDocumentParser (Python native)
    - PDF: PdfParserFactory selects PyMuPDF, LlamaParse, or LiteParse at runtime
{%- elif cookiecutter.use_llamaparse %}
    Supported file types:
    - TXT, MD: TextDocumentParser (Python native)
    - PDF, DOCX, PPTX, XLSX, images, and 130+ more: LlamaParseParser (cloud API)
{%- else %}
    Supported file types:
    - TXT, MD: TextDocumentParser (Python native)
    - DOCX: DocxDocumentParser (Python native)
    - PDF: PyMuPDFParser (local, tables, OCR fallback)
{%- endif %}
    """

    def __init__(self, settings: RAGSettings):
        """Initialize the document processor.

        Args:
            settings: RAG configuration settings.
        """
        self.settings = settings
        self.splitter = self._create_splitter(settings)

        # Always use Python native parser for plain text
        self.text_parser = TextDocumentParser()
        {%- if cookiecutter.use_all_pdf_parsers %}
        self.docx_parser = DocxDocumentParser()
        {%- if cookiecutter.enable_rag_image_description %}
        self.image_describer = self._init_image_describer(settings) if settings.enable_image_description else None
        {%- else %}
        self.image_describer = None
        {%- endif %}
        self.pdf_parser = PdfParserFactory.create(
            parser_name=settings.pdf_parser.method,
            settings=settings,
            image_describer=self.image_describer,
        )
        {%- elif cookiecutter.use_llamaparse %}
        # LlamaParse handles PDF, DOCX, PPTX, XLSX, images, and more
        self.llamaparse_parser = LlamaParseParser(api_key=settings.pdf_parser.api_key, tier=settings.pdf_parser.tier)
        {%- elif cookiecutter.use_liteparse %}
        # LiteParse handles PDFs with layout-aware text extraction
        self.liteparse_parser = LiteParseParser()
        self.docx_parser = DocxDocumentParser()
        {%- else %}
        self.docx_parser = DocxDocumentParser()
        {%- if cookiecutter.enable_rag_image_description %}
        # Image describer for LLM-based image descriptions and OCR fallback
        self.image_describer = self._init_image_describer(settings) if settings.enable_image_description else None
        self.pdf_parser = PyMuPDFParser(
            enable_ocr=settings.enable_ocr,
            image_describer=self.image_describer,
        )
        {%- else %}
        self.pdf_parser = PyMuPDFParser(enable_ocr=False)
        {%- endif %}
        {%- endif %}

    @staticmethod
    def _create_splitter(settings: RAGSettings) -> Any:
        """Create text splitter based on chunking strategy."""
        from langchain_text_splitters import (
            MarkdownHeaderTextSplitter,
            RecursiveCharacterTextSplitter,
        )

        strategy = settings.chunking_strategy

        if strategy == "markdown":
            # Split by markdown headers, then by size
            return MarkdownHeaderTextSplitter(
                headers_to_split_on=[
                    ("#", "h1"), ("##", "h2"), ("###", "h3"),
                ],
                strip_headers=False,
            )

        if strategy == "fixed":
            # Simple fixed-size chunks with no smart splitting
            return RecursiveCharacterTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
                length_function=len,
                separators=["\n"],
            )

        # Default: recursive (smart splitting by paragraphs, sentences, words)
        return RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

{%- if cookiecutter.enable_rag_image_description %}
    @staticmethod
    def _init_image_describer(settings: RAGSettings) -> Any:
        """Initialize the image describer using the configured AI framework."""
        from app.core.config import settings as app_settings

        model_name = getattr(app_settings, "RAG_IMAGE_DESCRIPTION_MODEL", None) or app_settings.AI_MODEL

{%- if cookiecutter.use_pydantic_ai %}
        from app.rag.image_describer import PydanticAIImageDescriber
        return PydanticAIImageDescriber(model_name=model_name)
{%- elif cookiecutter.use_langchain or cookiecutter.use_langgraph %}
        from app.rag.image_describer import LangChainImageDescriber
        return LangChainImageDescriber(model_name=model_name)
{%- elif cookiecutter.use_crewai %}
        from app.rag.image_describer import CrewAIImageDescriber
        return CrewAIImageDescriber(model_name=model_name)
{%- elif cookiecutter.use_deepagents %}
        from app.rag.image_describer import DeepAgentsImageDescriber
        return DeepAgentsImageDescriber(model_name=model_name)
{%- endif %}

    async def _describe_images(self, document: Document) -> None:
        """Generate text descriptions for all images in document pages."""
        for page in document.pages:
            if not page.images:
                continue
            for image in page.images:
                image.description = await self.image_describer.describe(
                    image.image_bytes, image.mime_type
                )
            img_descriptions = [
                f"[Image: {img.description}]"
                for img in page.images if img.description
            ]
            if img_descriptions:
                page.content = f"{page.content}\n\n{chr(10).join(img_descriptions)}"
{%- endif %}

    async def process_file(self, filepath: Path) -> Document:
        """Main entry point: filepath -> Document with chunks.

        Args:
            filepath: Path to the file to process.

        Returns:
            Document object with parsed pages and chunked content.

        Raises:
            ValueError: If the file type is not supported.
        """
        # Route to appropriate parser based on file extension
        if filepath.suffix in (".txt", ".md"):
            document = await self.text_parser.parse(filepath)
        {%- if cookiecutter.use_all_pdf_parsers %}
        elif filepath.suffix == ".docx":
            document = await self.docx_parser.parse(filepath)
        elif filepath.suffix == ".pdf":
            document = await self.pdf_parser.parse(filepath)
        {%- elif cookiecutter.use_llamaparse %}
        elif self.llamaparse_parser.is_extension_allowed(filepath):
            document = await self.llamaparse_parser.parse(filepath)
        {%- elif cookiecutter.use_liteparse %}
        elif filepath.suffix == ".docx":
            document = await self.docx_parser.parse(filepath)
        elif filepath.suffix == ".pdf":
            document = await self.liteparse_parser.parse(filepath)
        {%- else %}
        elif filepath.suffix == ".docx":
            document = await self.docx_parser.parse(filepath)
        elif filepath.suffix == ".pdf":
            document = await self.pdf_parser.parse(filepath)
        {%- endif %}
        else:
            raise ValueError(f"Unsupported file type: {filepath.suffix}")

{%- if cookiecutter.enable_rag_image_description %}
        # Describe images using LLM vision before chunking
        await self._describe_images(document)
{%- endif %}

        pages = document.pages

        chunked_pages: list[DocumentPageChunk] = []
        is_markdown_splitter = self.settings.chunking_strategy == "markdown"
        for page in pages:
            if is_markdown_splitter:
                # MarkdownHeaderTextSplitter returns Document objects
                md_docs = self.splitter.split_text(page.content)
                chunks = [doc.page_content for doc in md_docs]
            else:
                chunks = self.splitter.split_text(page.content)
            for chunk_num, chunk in enumerate(chunks):
                chunked_pages.append(DocumentPageChunk(
                    chunk_content=chunk,
                    chunk_num=chunk_num,
                    parent_doc_id=document.id,
                    **page.model_dump(
                        exclude={"parent_doc_id"}
                    )))

        # Add chunked pages to original document
        document.chunked_pages = chunked_pages
        return document

{%- endif %}
