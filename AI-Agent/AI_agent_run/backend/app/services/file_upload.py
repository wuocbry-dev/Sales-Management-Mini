"""File upload service (SQLite sync).

Contains business logic for file validation, content parsing, and chat file
creation. Moves parsing helpers and file classification out of the route layer.
"""

import csv
import html
import io
import json
import logging
import re
import zipfile
from typing import Any
from xml.etree import ElementTree

from sqlalchemy.orm import Session

from app.db.models.chat_file import ChatFile
from app.services.file_storage import (
    ALLOWED_MIME_TYPES,
    MAX_UPLOAD_SIZE,
    classify_file,
)

logger = logging.getLogger(__name__)


class FileUploadService:
    """Service for file upload validation, parsing, and persistence."""

    ALLOWED_MIME_TYPES = ALLOWED_MIME_TYPES
    MAX_UPLOAD_SIZE = MAX_UPLOAD_SIZE

    def __init__(self, db: Session):
        self.db = db

    @staticmethod
    def validate_upload(content_type: str | None, size: int) -> tuple[bool, str | None]:
        """Validate file type and size.

        Returns:
            Tuple of (is_valid, error_message).
        """
        if content_type not in ALLOWED_MIME_TYPES:
            return False, f"File type '{content_type}' is not supported."
        if size > MAX_UPLOAD_SIZE:
            return False, f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024 * 1024)}MB."
        return True, None

    @staticmethod
    def classify_file(mime_type: str, filename: str) -> str:
        """Classify file type based on MIME type and extension."""
        return classify_file(mime_type, filename)

    def parse_content(
        self,
        data: bytes,
        file_type: str,
        mime_type: str = "",
    ) -> str | None:
        """Parse file content based on file type.

        Returns extracted text content or None if parsing fails.
        """
        if file_type == "text":
            return self._parse_text_content(data, mime_type)
        elif file_type == "pdf":
            return self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
        elif file_type == "spreadsheet":
            return self._parse_spreadsheet_content(data)
        return None

    @staticmethod
    def _parse_text_content(data: bytes, mime_type: str) -> str | None:
        """Extract text content from text-based files."""
        for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1258", "cp1252", "latin-1"):
            try:
                text = data.decode(encoding)
                break
            except (UnicodeDecodeError, ValueError):
                continue
        else:
            return None

        text = text.replace("\x00", "").strip()
        if not text:
            return None

        lowered_mime = mime_type.lower()
        if "json" in lowered_mime:
            try:
                return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                return text
        if "csv" in lowered_mime:
            return FileUploadService._normalize_csv_text(text)
        if "html" in lowered_mime or "<html" in text[:500].lower():
            return FileUploadService._html_to_text(text)
        return text

    @staticmethod
    def _normalize_csv_text(text: str) -> str:
        """Normalize CSV-ish text into a readable markdown-style table."""
        try:
            sample = text[:4096]
            dialect = csv.Sniffer().sniff(sample)
            rows = list(csv.reader(io.StringIO(text), dialect))
            return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
        except Exception:
            return text

    @staticmethod
    def _html_to_text(text: str) -> str:
        """Best-effort HTML to plain text without extra dependencies."""
        cleaned = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
        cleaned = re.sub(r"(?i)<br\s*/?>", "\n", cleaned)
        cleaned = re.sub(r"(?i)</p\s*>", "\n", cleaned)
        cleaned = re.sub(r"<[^>]+>", " ", cleaned)
        cleaned = html.unescape(cleaned)
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n\s*\n+", "\n\n", cleaned)
        return cleaned.strip() or text

    @staticmethod
    def _parse_pdf_content(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call,unused-ignore]
            texts = []
            for page in doc:
                blocks = page.get_text("blocks")
                for b in blocks:
                    if b[6] == 0:
                        text = b[4].strip()
                        if text:
                            texts.append(text)
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for table in tables.tables:
                            df = table.to_pandas()
                            if not df.empty:
                                texts.append(df.to_markdown(index=False))
                except Exception:
                    pass
            doc.close()
            return "\n\n".join(texts) if texts else None
        except Exception as e:
            logger.warning(f"PDF parsing failed: {e}")
            return None

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io

            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts) if parts else None
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None

    @staticmethod
    def _parse_spreadsheet_content(data: bytes) -> str | None:
        """Extract readable text from XLSX files using only stdlib zip/xml."""
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                shared_strings = FileUploadService._read_xlsx_shared_strings(archive)
                sheet_names = sorted(
                    name
                    for name in archive.namelist()
                    if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
                )
                sections: list[str] = []
                for sheet_index, sheet_name in enumerate(sheet_names, start=1):
                    rows = FileUploadService._read_xlsx_sheet(
                        archive.read(sheet_name), shared_strings
                    )
                    if rows:
                        sections.append(
                            "\n".join(
                                [f"Sheet {sheet_index}"]
                                + [" | ".join(cell for cell in row if cell) for row in rows]
                            )
                        )
                return "\n\n".join(sections) if sections else None
        except Exception as e:
            logger.warning(f"Spreadsheet parsing failed: {e}")
            return None

    @staticmethod
    def _read_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
        try:
            root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
        except KeyError:
            return []
        namespace = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        values: list[str] = []
        for item in root.findall("x:si", namespace):
            texts = [node.text or "" for node in item.findall(".//x:t", namespace)]
            values.append("".join(texts))
        return values

    @staticmethod
    def _read_xlsx_sheet(sheet_xml: bytes, shared_strings: list[str]) -> list[list[str]]:
        root = ElementTree.fromstring(sheet_xml)
        namespace = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        rows: list[list[str]] = []
        for row in root.findall(".//x:sheetData/x:row", namespace):
            values: list[str] = []
            for cell in row.findall("x:c", namespace):
                cell_type = cell.attrib.get("t")
                value_node = cell.find("x:v", namespace)
                inline_text = cell.find(".//x:is/x:t", namespace)
                raw = value_node.text if value_node is not None else inline_text.text if inline_text is not None else ""
                if cell_type == "s" and raw:
                    try:
                        values.append(shared_strings[int(raw)])
                    except (ValueError, IndexError):
                        values.append(raw)
                else:
                    values.append(raw or "")
            if any(value.strip() for value in values):
                rows.append(values)
        return rows

    def get_user_file(self, file_id: Any, user_id: Any) -> ChatFile:
        """Get a file by ID, verifying ownership.

        Raises:
            NotFoundError: If file does not exist or user has no access.
        """
        from app.core.exceptions import NotFoundError
        from app.repositories import chat_file as chat_file_repo

        chat_file = chat_file_repo.get_by_id(self.db, file_id)
        if not chat_file or str(chat_file.user_id) != str(user_id):
            raise NotFoundError(message="File not found")
        return chat_file

    def create_chat_file(
        self,
        *,
        user_id: Any,
        filename: str,
        mime_type: str,
        size: int,
        storage_path: str,
        file_type: str,
        parsed_content: str | None = None,
    ) -> ChatFile:
        """Create a chat file record in the database."""
        chat_file = ChatFile(
            user_id=user_id,
            filename=filename,
            mime_type=mime_type,
            size=size,
            storage_path=storage_path,
            file_type=file_type,
            parsed_content=parsed_content,
        )
        self.db.add(chat_file)
        self.db.flush()
        self.db.commit()
        self.db.refresh(chat_file)
        return chat_file
